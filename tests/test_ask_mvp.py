import io
import sqlite3

import pandas as pd
import pytest
from PIL import Image
from docx import Document
from pypdf import PdfWriter
from werkzeug.datastructures import FileStorage

from app import create_app
from app.database.migrations import upgrade
from app.workspace.repositories.ask_repository import AskRepository
from app.workspace.services.ask_analysis_engine import AskAnalysisEngine
from app.workspace.services.ask_artifact_export_service import (
    AskArtifactExportService,
)
from app.workspace.services.ask_execution_service import AskExecutionService
from app.workspace.services.ask_conversation_service import (
    AskConversationService,
)
from app.workspace.services.ask_file_service import AskFileService
from app.workspace.services.ask_openai_service import (
    AskOpenAIError, AskOpenAIService,
)
from app.workspace.services.ask_preparation_service import AskPreparationService


@pytest.fixture
def ask_database(tmp_path, monkeypatch):
    path = tmp_path / "ask.db"
    monkeypatch.setattr("app.database.connection.DB_PATH", path)
    monkeypatch.setattr(AskFileService, "STORAGE_ROOT", tmp_path / "ask-files")
    upgrade()
    with sqlite3.connect(path) as connection:
        connection.execute(
            """INSERT INTO ws_customers(id,name,erp_customer_id)
            VALUES(1,'CARTONES AMÉRICA','900100200')"""
        )
        connection.execute(
            """INSERT INTO ws_customers(id,name,erp_customer_id)
            VALUES(2,'SERMOTOR INGENIERIA S.A.S.','900611187')"""
        )
        connection.execute(
            """INSERT INTO ws_projects(
                id,customer_id,name,status,objective,sales_rep
            ) VALUES(1,1,'Prueba INA','waiting_customer','Prueba','Asesor')"""
        )
        connection.execute(
            """INSERT INTO ws_project_brands(project_id,brand)
            VALUES(1,'INA')"""
        )
        connection.execute(
            """INSERT INTO ws_project_brands(project_id,brand)
            VALUES(1,'SKF')"""
        )
    return path


def _spreadsheet(rows=None, filename="evidencia.xlsx"):
    stream = io.BytesIO()
    pd.DataFrame(rows or [
        {"Cuenta": "Cliente A", "Indicador": 18.4, "Prioridad": "Alta"},
        {"Cuenta": "Cliente B", "Indicador": 7.2, "Prioridad": "Media"},
    ]).to_excel(stream, index=False)
    stream.seek(0)
    return FileStorage(stream=stream, filename=filename)


def _knowledge_synthesis(_knowledge):
    return {
        "title": "Conocimiento consolidado",
        "confidence": "Alta",
        "unresolved_questions": [],
        "sections": [{
            "title": "Hallazgos provisionales",
            "type": "list",
            "content": ["La evidencia fue procesada."],
        }],
    }


def test_planner_has_no_inventory_assumptions_or_column_contract(
    ask_database,
):
    analysis_id = AskPreparationService.create(
        "Evalúe la rentabilidad y ayúdeme a preparar la decisión", 2
    )
    AskFileService.upload(analysis_id, _spreadsheet())

    page = AskPreparationService.refresh(analysis_id)
    analysis = page["analysis"]

    assert analysis["assumptions"] == {}
    assert analysis["mappings"] == {}
    assert analysis["plan"]["mode"] == "investigation"
    assert analysis["plan"]["capabilities"] == ["uploaded_evidence"]
    serialized = str(analysis["plan"]).casefold()
    assert "coverage" not in serialized
    assert "historical_months" not in serialized
    assert "proposed_quantity" not in serialized


def test_investigation_produces_knowledge_not_automatic_artifacts(
    ask_database, monkeypatch,
):
    analysis_id = AskPreparationService.create(
        "Compare las cuentas del archivo y encuentre diferencias", 2
    )
    AskFileService.upload(analysis_id, _spreadsheet())
    AskPreparationService.refresh(analysis_id)
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)

    completed = AskExecutionService.execute(analysis_id)
    knowledge = completed["evidence"]

    assert completed["status"] == "completed"
    assert knowledge["goal"].startswith("Compare")
    assert knowledge["summary"]["facts"] == 1
    assert knowledge["working_dataset"][0]["Cuenta"] == "Cliente A"
    assert knowledge["trace"][-1]["step"] == (
        "Conocimiento estructurado consolidado"
    )
    assert AskRepository.list_artifacts(analysis_id) == []


def test_follow_up_creates_version_and_requested_dynamic_artifact(
    ask_database, monkeypatch,
):
    analysis_id = AskPreparationService.create(
        "Revise la evidencia para apoyar una decisión", 2
    )
    AskFileService.upload(analysis_id, _spreadsheet())
    AskPreparationService.refresh(analysis_id)
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)
    AskExecutionService.execute(analysis_id)

    new_id = AskPreparationService.continue_investigation(
        analysis_id, 2,
        "Muéstreme qué se aprueba, qué no y prepare la reunión con el equipo.",
    )
    version = AskRepository.get(new_id)
    assert version["version"] == 2
    assert version["evidence"]["facts"]
    assert version["plan"]["mode"] == "deliverable"

    monkeypatch.setattr(
        AskOpenAIService,
        "specify_artifacts",
        lambda knowledge, instruction: {
            "artifacts": [{
                "key": "decision-team",
                "type": "document",
                "title": "Decisiones para el equipo",
                "blocks": [{
                    "type": "list",
                    "title": "Decisiones",
                    "content": ["Revisar Cliente A con el equipo."],
                }],
                "metadata": {},
            }],
        },
    )
    AskExecutionService.execute(new_id)
    artifacts = AskRepository.list_artifacts(new_id)

    assert artifacts[0]["key"] == "decision-team"
    assert artifacts[0]["blocks"][0]["title"] == "Decisiones"
    assert AskRepository.list_artifacts(analysis_id) == []


def test_conversation_automatically_investigates_and_builds_deliverable(
    ask_database, monkeypatch,
):
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)
    analysis_id = AskConversationService.start(
        "Revise esta decisión", [_spreadsheet()], 2
    )

    assert AskRepository.get(analysis_id)["status"] == "completed"
    assert AskRepository.list_messages(analysis_id)[-1][
        "clarification_type"
    ] == "knowledge_ready"
    monkeypatch.setattr(
        AskOpenAIService,
        "specify_artifacts",
        lambda knowledge, instruction: {
            "artifacts": [{
                "key": "resumen-equipo", "type": "document",
                "title": "Resumen para el equipo",
                "blocks": [{
                    "type": "text", "title": "Decisión",
                    "content": "Revisar con el equipo.",
                }],
                "metadata": {},
            }],
        },
    )

    new_id = AskConversationService.respond(
        analysis_id, 2, "Prepare un resumen para el equipo.", []
    )

    assert AskRepository.get(new_id)["status"] == "completed"
    assert AskRepository.list_artifacts(new_id)[0]["key"] == "resumen-equipo"
    assert AskRepository.list_messages(new_id)[-1][
        "clarification_type"
    ] == "artifacts_ready"


def test_new_evidence_after_execution_versions_without_mutating_history(
    ask_database, monkeypatch,
):
    analysis_id = AskPreparationService.create("Revise estas fuentes", 2)
    AskFileService.upload(analysis_id, _spreadsheet())
    AskPreparationService.refresh(analysis_id)
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)
    AskExecutionService.execute(analysis_id)
    client = create_app({
        "TESTING": True, "TEST_AUTH_BYPASS": True,
        "TEST_AUTH_USER_ID": 2,
    }, run_migrations=False).test_client()

    response = client.post(
        f"/ask/analysis/{analysis_id}/files",
        data={"files": (io.BytesIO(b"Nueva evidencia"), "nota.txt")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 302
    new_id = int(response.location.rstrip("/").rsplit("/", 1)[-1])
    assert AskRepository.get(new_id)["version"] == 2
    assert len(AskRepository.list_files(analysis_id)) == 1
    assert len(AskRepository.list_files(new_id)) == 2


def test_spreadsheet_profiles_every_sheet_without_expected_columns(
    ask_database,
):
    stream = io.BytesIO()
    with pd.ExcelWriter(stream) as writer:
        pd.DataFrame({"Producto": ["A"], "Costo": [100]}).to_excel(
            writer, sheet_name="Costos", index=False
        )
        pd.DataFrame({"Zona": ["Cali"], "Margen": [20]}).to_excel(
            writer, sheet_name="Márgenes", index=False
        )
    stream.seek(0)
    analysis_id = AskPreparationService.create("Compare las fuentes", 2)
    AskFileService.upload(
        analysis_id, FileStorage(stream, filename="comparacion.xlsx")
    )

    inspection = AskRepository.list_files(analysis_id)[0]["inspection"]

    assert inspection["worksheets"] == ["Costos", "Márgenes"]
    assert len(inspection["tables"]) == 2
    assert "proposed_mappings" not in inspection
    assert inspection["tables"][1]["columns"][0]["name"] == "Zona"


def test_generic_file_evidence_supports_documents_and_images(ask_database):
    analysis_id = AskPreparationService.create("Investigue los anexos", 2)
    docx_stream = io.BytesIO()
    document = Document()
    document.add_heading("Hallazgo", level=1)
    document.add_paragraph("Existe evidencia pendiente de validación.")
    document.save(docx_stream)
    docx_stream.seek(0)
    pdf_stream = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(pdf_stream)
    pdf_stream.seek(0)
    image_stream = io.BytesIO()
    Image.new("RGB", (32, 24), "white").save(image_stream, format="PNG")
    image_stream.seek(0)

    for upload in (
        FileStorage(io.BytesIO(b"Nota"), filename="nota.txt"),
        FileStorage(docx_stream, filename="soporte.docx"),
        FileStorage(pdf_stream, filename="anexo.pdf"),
        FileStorage(image_stream, filename="foto.png"),
    ):
        AskFileService.upload(analysis_id, upload)

    files = AskRepository.list_files(analysis_id)
    assert [file["inspection"]["kind"] for file in files] == [
        "text", "document", "pdf", "image",
    ]
    assert all(file["processing_status"] == "processed" for file in files)


def test_workspace_is_conversation_first_and_runs_without_workflow_ui(
    ask_database, monkeypatch,
):
    analysis_id = AskPreparationService.create(
        "Analice el desempeño comercial", 2
    )
    AskPreparationService.refresh(analysis_id)
    app = create_app({
        "TESTING": True, "TEST_AUTH_BYPASS": True,
        "TEST_AUTH_USER_ID": 2,
    }, run_migrations=False)
    client = app.test_client()
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)

    preparation = client.get(f"/ask/analysis/{analysis_id}")
    assert preparation.status_code == 200
    assert b"Ask \xc2\xb7 Analista comercial" in preparation.data
    assert b"Enviar" in preparation.data
    assert b"Continuar" not in preparation.data
    assert b"Realizar investigaci" not in preparation.data
    assert b"Plan de investigaci" not in preparation.data
    assert b"Historical months" not in preparation.data
    assert b"Coverage" not in preparation.data
    assert AskRepository.get(analysis_id)["status"] == "completed"
    assert b"La investigaci" in preparation.data


def test_openai_failure_preserves_structured_knowledge(
    ask_database, monkeypatch,
):
    analysis_id = AskPreparationService.create("Revise la fuente", 2)
    AskFileService.upload(analysis_id, _spreadsheet())
    AskPreparationService.refresh(analysis_id)

    def fail(_knowledge):
        raise AskOpenAIError("OpenAI no disponible.")

    monkeypatch.setattr(AskOpenAIService, "generate", fail)
    with pytest.raises(AskOpenAIError):
        AskExecutionService.execute(analysis_id)

    failed = AskRepository.get(analysis_id)
    assert failed["status"] == "failed"
    assert failed["evidence"]["facts"]
    assert failed["evidence"]["working_dataset"]


def test_artifact_contract_is_metadata_driven():
    knowledge = {
        "confidence": "Alta",
        "facts": [{"statement": "Hecho validado"}],
        "findings": [], "risks": [], "opportunities": [],
        "recommendations": [], "pending_investigations": [],
        "supporting_evidence": [],
        "working_dataset": [{
            "Cliente": "A", "Rentabilidad": 18.4, "Decisión": "Revisar",
        }],
    }

    artifacts = AskAnalysisEngine.build_artifacts(knowledge, {
        "title": "Rentabilidad",
        "sections": [{
            "title": "Resultado", "type": "text",
            "content": "Hecho validado",
        }],
        "include_dataset": True,
    })

    assert [artifact["type"] for artifact in artifacts] == [
        "document", "dataset",
    ]
    schema = artifacts[1]["blocks"][0]["schema"]
    assert [column["name"] for column in schema] == [
        "Cliente", "Rentabilidad", "Decisión",
    ]


def test_table_artifact_exports_excel_and_csv(ask_database):
    analysis_id = AskPreparationService.create("Prepare una tabla", 2)
    AskRepository.replace_artifacts(analysis_id, [{
        "key": "decisiones", "type": "dataset", "title": "Decisiones",
        "blocks": [{
            "type": "table", "title": "Resultado",
            "schema": [], "rows": [{"Cliente": "A", "Decisión": "Revisar"}],
        }],
        "metadata": {},
    }])

    excel, excel_type, excel_name = AskArtifactExportService.export(
        analysis_id, "decisiones", "xlsx"
    )
    csv, csv_type, csv_name = AskArtifactExportService.export(
        analysis_id, "decisiones", "csv"
    )

    assert excel.read(2) == b"PK"
    assert excel_type.endswith("spreadsheetml.sheet")
    assert excel_name == "decisiones.xlsx"
    assert "Cliente" in csv.read().decode("utf-8-sig")
    assert csv_type == "text/csv"
    assert csv_name == "decisiones.csv"


def test_brand_detection_does_not_find_ina_inside_consignacion(ask_database):
    from app.workspace.repositories.ask_context_repository import (
        AskContextRepository,
    )

    assert AskContextRepository.brand_candidates(["consignacion"]) == []
    assert AskContextRepository.brand_candidates(["SKF"]) == ["SKF"]


def test_explicit_customer_selection_avoids_repeated_questions(
    ask_database, monkeypatch,
):
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)
    monkeypatch.setattr(
        AskOpenAIService, "plan_investigation",
        lambda context: {
            "questions": [
                "¿Existen políticas internas para definir cantidades?"
            ],
            "missing_evidence": [
                "Criterios internos de consignación"
            ],
            "capabilities": context["available_capabilities"],
        },
    )

    analysis_id = AskConversationService.start(
        "Analice la propuesta de bodega en consignación",
        [_spreadsheet()], 2, customer_id=2,
    )
    analysis = AskRepository.get(analysis_id)

    assert analysis["customer_id"] == 2
    assert analysis["status"] == "completed"
    assert analysis["blocking_reasons"] == []
    assert analysis["context"]["customer"]["name"] == (
        "SERMOTOR INGENIERIA S.A.S."
    )


def test_conversation_composer_contains_customer_autocomplete(
    ask_database, monkeypatch,
):
    monkeypatch.setattr(AskOpenAIService, "generate", _knowledge_synthesis)
    analysis_id = AskConversationService.start(
        "Analice esta decisión", [], 2
    )
    client = create_app({
        "TESTING": True, "TEST_AUTH_BYPASS": True,
        "TEST_AUTH_USER_ID": 2,
    }, run_migrations=False).test_client()

    response = client.get(f"/ask/analysis/{analysis_id}")

    assert response.status_code == 200
    assert b'id="ask-customer-search"' in response.data
    assert b"Cliente (opcional)" in response.data
