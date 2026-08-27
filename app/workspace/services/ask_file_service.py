import hashlib
import mimetypes
import uuid
from pathlib import Path
from typing import Any

import pandas as pd
from PIL import ExifTags, Image, UnidentifiedImageError
from docx import Document
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from app.database.transaction import transactional
from app.storage import upload_path
from app.workspace.connectors.workbook_reader import (
    WorkbookReadError, WorkbookReader,
)
from app.workspace.repositories.ask_repository import AskRepository


class AskFileError(ValueError):
    pass


class AskFileService:
    """Stores files and turns them into bounded, format-neutral evidence."""

    STORAGE_ROOT = upload_path("ask")
    ALLOWED_EXTENSIONS = {
        ".xlsx", ".xls", ".csv", ".pdf", ".docx", ".txt",
        ".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff",
    }
    MAX_BYTES = 20 * 1024 * 1024
    MAX_ROWS = 20_000
    MAX_TEXT_EXCERPT = 8_000
    IMAGE_EXTENSIONS = {
        ".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff",
    }

    @classmethod
    @transactional
    def upload(cls, analysis_id: int, upload: FileStorage) -> int:
        if not upload or not upload.filename:
            raise AskFileError("Seleccione un archivo.")
        extension = Path(upload.filename).suffix.casefold()
        if extension not in cls.ALLOWED_EXTENSIONS:
            raise AskFileError(
                "Formato no soportado. Use XLSX, XLS, CSV, PDF, DOCX, TXT "
                "o una imagen común."
            )
        cls.STORAGE_ROOT.mkdir(parents=True, exist_ok=True)
        safe_name = secure_filename(upload.filename) or f"archivo{extension}"
        stored_name = f"{uuid.uuid4().hex}{extension}"
        path = cls.STORAGE_ROOT / stored_name
        upload.save(path)
        size = path.stat().st_size
        if size > cls.MAX_BYTES:
            path.unlink(missing_ok=True)
            raise AskFileError("El archivo supera el límite de 20 MB.")
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        try:
            inspection = cls.inspect(path, extension)
            status, error = "processed", None
        except (
            AskFileError, WorkbookReadError, ValueError, OSError,
            UnidentifiedImageError,
        ) as exception:
            inspection, status, error = {}, "failed", str(exception)
        return AskRepository.add_file(analysis_id, {
            "original_filename": upload.filename,
            "stored_filename": stored_name, "stored_path": str(path),
            "file_extension": extension,
            "mime_type": upload.mimetype or mimetypes.guess_type(safe_name)[0],
            "file_size_bytes": size, "file_hash": digest,
            "processing_status": status, "inspection": inspection,
            "error_message": error,
        })

    @classmethod
    @transactional
    def remove(cls, analysis_id: int, file_id: int) -> None:
        file = AskRepository.get_file(file_id)
        if not file or file["analysis_id"] != analysis_id:
            raise AskFileError("El archivo no pertenece a este análisis.")
        AskRepository.delete_file(file_id)

    @classmethod
    def inspect(cls, path: Path, extension: str) -> dict[str, Any]:
        if extension == ".txt":
            return cls._inspect_text(path)
        if extension == ".pdf":
            return cls._inspect_pdf(path)
        if extension == ".docx":
            return cls._inspect_docx(path)
        if extension in cls.IMAGE_EXTENSIONS:
            return cls._inspect_image(path)
        return cls._inspect_tabular(path, extension)

    @classmethod
    def structured_evidence(cls, file_record: dict[str, Any]) -> dict[str, Any]:
        """Return the bounded evidence contract supplied to the engine/LLM."""
        inspection = file_record.get("inspection") or {}
        return {
            "file_id": file_record["id"],
            "filename": file_record["original_filename"],
            "file_type": inspection.get("kind")
            or file_record["file_extension"].lstrip("."),
            "mime_type": file_record.get("mime_type"),
            "size_bytes": file_record["file_size_bytes"],
            "sha256": file_record["file_hash"],
            "processing_status": file_record["processing_status"],
            "structure": inspection.get("structure", {}),
            "evidence": inspection.get("evidence", {}),
            "warnings": inspection.get("warnings", []),
        }

    @classmethod
    def read_dataset(
        cls, file_record: dict[str, Any], worksheet: str | None = None
    ) -> pd.DataFrame:
        path = Path(file_record["stored_path"])
        extension = file_record["file_extension"]
        if extension == ".csv":
            return cls._read_csv(path, cls.MAX_ROWS)
        if extension in {".xlsx", ".xls"}:
            result = WorkbookReader.read(path, worksheet, cls.MAX_ROWS)
            columns = cls._unique_columns(result["header"])
            return pd.DataFrame(
                [row[:len(columns)] for _, row in result["rows"]],
                columns=columns,
            )
        raise AskFileError("El archivo no contiene datos tabulares.")

    @classmethod
    def _inspect_tabular(cls, path: Path, extension: str) -> dict[str, Any]:
        if extension == ".csv":
            frame = cls._read_csv(path, cls.MAX_ROWS + 1)
            if len(frame) > cls.MAX_ROWS:
                raise AskFileError(
                    f"El archivo supera el límite de {cls.MAX_ROWS} filas."
                )
            tables = [cls._profile_frame(frame, "CSV")]
            selected = "CSV"
        else:
            first = WorkbookReader.read(path, None, cls.MAX_ROWS)
            selected = first["selected_worksheet"]
            tables = []
            for worksheet in first.get("worksheets", []) or [selected]:
                workbook = (
                    first if worksheet == selected
                    else WorkbookReader.read(path, worksheet, cls.MAX_ROWS)
                )
                columns = cls._unique_columns(workbook["header"])
                frame = pd.DataFrame(
                    [row[:len(columns)] for _, row in workbook["rows"]],
                    columns=columns,
                )
                tables.append(cls._profile_frame(frame, worksheet))
        primary = next(
            (table for table in tables if table["name"] == selected),
            tables[0] if tables else cls._profile_frame(pd.DataFrame(), selected),
        )
        return {
            "kind": "spreadsheet",
            "worksheets": [table["name"] for table in tables],
            "selected_worksheet": selected,
            # Backwards-compatible conveniences; interpretation is not stored here.
            "columns": [column["name"] for column in primary["columns"]],
            "row_count": primary["row_count"],
            "numeric_columns": [
                column["name"] for column in primary["columns"]
                if column["inferred_type"] == "number"
            ],
            "sample_rows": primary["sample_rows"],
            "tables": tables,
            "structure": {
                "sheet_count": len(tables),
                "sheets": [{
                    "name": table["name"],
                    "rows": table["row_count"],
                    "columns": len(table["columns"]),
                    "column_names": [
                        column["name"] for column in table["columns"]
                    ],
                } for table in tables],
            },
            "evidence": {
                "summary": (
                    f"{len(tables)} hoja(s) o tabla(s); "
                    f"{sum(table['row_count'] for table in tables)} filas."
                ),
                "tables": tables,
            },
            "warnings": [],
        }

    @classmethod
    def _inspect_text(cls, path: Path) -> dict[str, Any]:
        raw = path.read_bytes()
        if b"\x00" in raw:
            raise AskFileError("El archivo TXT no contiene texto válido.")
        text = raw.decode("utf-8-sig", errors="replace")
        if text and text.count("\ufffd") / len(text) > .01:
            raise AskFileError("El archivo TXT no contiene texto válido.")
        lines = text.splitlines()
        sections = cls._text_sections(text)
        return {
            "kind": "text",
            "character_count": len(text),
            "sample": text[:1000],
            "structure": {
                "line_count": len(lines),
                "word_count": len(text.split()),
                "section_count": len(sections),
            },
            "evidence": {
                "summary": f"Documento de texto con {len(lines)} líneas.",
                "excerpt": text[:cls.MAX_TEXT_EXCERPT],
                "sections": sections,
            },
            "warnings": (
                ["El texto fue truncado para razonamiento."]
                if len(text) > cls.MAX_TEXT_EXCERPT else []
            ),
        }

    @classmethod
    def _inspect_pdf(cls, path: Path) -> dict[str, Any]:
        try:
            reader = PdfReader(path)
            if reader.is_encrypted and not reader.decrypt(""):
                raise AskFileError("El PDF está protegido y no puede leerse.")
            pages = []
            budget = cls.MAX_TEXT_EXCERPT
            for number, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                excerpt = text[:max(0, min(1200, budget))]
                budget -= len(excerpt)
                pages.append({
                    "page": number, "character_count": len(text),
                    "excerpt": excerpt,
                })
        except AskFileError:
            raise
        except Exception as error:
            raise AskFileError("El PDF está corrupto o no es válido.") from error
        no_text = not any(page["character_count"] for page in pages)
        return {
            "kind": "pdf",
            "structure": {
                "page_count": len(pages),
                "metadata": {
                    str(key).lstrip("/"): str(value)
                    for key, value in (reader.metadata or {}).items()
                    if value is not None
                },
            },
            "evidence": {
                "summary": f"PDF con {len(pages)} página(s).",
                "pages": pages,
            },
            "warnings": (
                ["No se detectó texto extraíble; el PDF puede ser escaneado."]
                if no_text else []
            ),
        }

    @classmethod
    def _inspect_docx(cls, path: Path) -> dict[str, Any]:
        try:
            document = Document(path)
        except Exception as error:
            raise AskFileError("El DOCX está corrupto o no es válido.") from error
        paragraphs = [
            paragraph.text.strip() for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        sections = []
        current = {"title": "Documento", "excerpt": ""}
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                if current["excerpt"]:
                    sections.append(current)
                current = {"title": text, "excerpt": ""}
            elif len(current["excerpt"]) < 1200:
                current["excerpt"] = (
                    current["excerpt"] + "\n" + text
                ).strip()[:1200]
        if current["excerpt"] or current["title"] != "Documento":
            sections.append(current)
        tables = []
        for index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            frame = pd.DataFrame(
                rows[1:], columns=cls._unique_columns(rows[0])
            ) if rows else pd.DataFrame()
            tables.append(cls._profile_frame(frame, f"Tabla {index}"))
        excerpt = "\n".join(paragraphs)[:cls.MAX_TEXT_EXCERPT]
        return {
            "kind": "document",
            "structure": {
                "paragraph_count": len(paragraphs),
                "section_count": len(sections),
                "table_count": len(tables),
            },
            "evidence": {
                "summary": (
                    f"DOCX con {len(paragraphs)} párrafo(s) y "
                    f"{len(tables)} tabla(s)."
                ),
                "excerpt": excerpt, "sections": sections[:30],
                "tables": tables,
            },
            "warnings": (
                ["El texto fue truncado para razonamiento."]
                if len("\n".join(paragraphs)) > cls.MAX_TEXT_EXCERPT else []
            ),
        }

    @classmethod
    def _inspect_image(cls, path: Path) -> dict[str, Any]:
        try:
            with Image.open(path) as image:
                image.verify()
            with Image.open(path) as image:
                exif = {
                    ExifTags.TAGS.get(key, str(key)): str(value)
                    for key, value in image.getexif().items()
                    if ExifTags.TAGS.get(key, str(key)) in {
                        "DateTime", "DateTimeOriginal", "Make", "Model",
                        "Software", "ImageDescription",
                    }
                }
                structure = {
                    "format": image.format,
                    "width": image.width, "height": image.height,
                    "mode": image.mode,
                    "frame_count": getattr(image, "n_frames", 1),
                    "metadata": exif,
                }
        except (UnidentifiedImageError, OSError) as error:
            raise AskFileError("La imagen está corrupta o no es válida.") from error
        return {
            "kind": "image",
            "structure": structure,
            "evidence": {
                "summary": (
                    f"Imagen {structure['format']} de "
                    f"{structure['width']} × {structure['height']} px."
                ),
                "metadata": structure["metadata"],
            },
            "warnings": [
                "No se ejecutó OCR local; solo se extrajeron estructura y metadatos."
            ],
        }

    @classmethod
    def _profile_frame(
        cls, frame: pd.DataFrame, name: str | None
    ) -> dict[str, Any]:
        frame = frame.copy()
        frame.columns = cls._unique_columns(list(frame.columns))
        columns = []
        for column in frame.columns:
            series = frame[column]
            non_null = series.dropna()
            numeric = pd.to_numeric(non_null, errors="coerce")
            inferred = (
                "number"
                if len(non_null) and numeric.notna().mean() >= .8
                else "date"
                if len(non_null) and pd.to_datetime(
                    non_null, errors="coerce", format="mixed"
                ).notna().mean() >= .8
                else "text"
            )
            columns.append({
                "name": str(column), "inferred_type": inferred,
                "non_null_count": int(non_null.size),
                "unique_count": int(non_null.astype(str).nunique()),
                "sample_values": [
                    cls._json_value(value) for value in non_null.iloc[:5]
                ],
            })
        return {
            "name": name or "Tabla",
            "row_count": len(frame),
            "columns": columns,
            "sample_rows": [
                {
                    str(column): cls._json_value(value)
                    for column, value in row.items()
                }
                for row in frame.head(10).to_dict(orient="records")
            ],
        }

    @staticmethod
    def _read_csv(path: Path, limit: int) -> pd.DataFrame:
        try:
            return pd.read_csv(path, nrows=limit)
        except UnicodeDecodeError:
            return pd.read_csv(path, nrows=limit, encoding="latin-1")
        except Exception as error:
            raise AskFileError("El CSV está corrupto o no es válido.") from error

    @staticmethod
    def _text_sections(text: str) -> list[dict[str, str]]:
        chunks = [chunk.strip() for chunk in text.split("\n\n") if chunk.strip()]
        return [
            {
                "title": chunk.splitlines()[0][:120],
                "excerpt": chunk[:1200],
            }
            for chunk in chunks[:30]
        ]

    @staticmethod
    def _unique_columns(header: list[Any]) -> list[str]:
        columns = []
        occurrences: dict[str, int] = {}
        for index, value in enumerate(header, start=1):
            base = str(value or "").strip() or f"Columna {index}"
            occurrences[base] = occurrences.get(base, 0) + 1
            suffix = occurrences[base]
            columns.append(base if suffix == 1 else f"{base} ({suffix})")
        return columns

    @staticmethod
    def _json_value(value: Any) -> Any:
        if value is None or pd.isna(value):
            return None
        if hasattr(value, "isoformat"):
            return value.isoformat()
        if isinstance(value, (str, int, float, bool)):
            return value
        return str(value)
